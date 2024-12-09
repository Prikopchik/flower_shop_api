from reportlab.pdfgen import canvas
from django.http import HttpResponse
from .models import Flower, Order, Customer
from datetime import datetime
from docx import Document

def generate_report(request):
    # HTTP response для PDF
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="flower_shop_report_{datetime.now().date()}.pdf"'

    # Создаем PDF с ReportLab
    p = canvas.Canvas(response)
    p.setTitle("Flower Shop Report")
    
    # Титульный лист
    p.setFont("Helvetica-Bold", 20)
    p.drawString(200, 800, "Отчет по цветочному магазину")
    p.setFont("Helvetica", 12)
    p.drawString(200, 770, f"Дата создания: {datetime.now().date()}")
    p.line(50, 760, 550, 760)

    # Описание сервиса
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, 730, "Краткое описание сервиса:")
    p.setFont("Helvetica", 12)
    p.drawString(50, 710, "Сервис предоставляет REST API для работы с цветами, клиентами и заказами.")

    # Таблица товаров
    p.setFont("Helvetica-Bold", 14)
    p.drawString(50, 680, "Таблица цветов:")
    y = 660
    p.setFont("Helvetica", 12)
    p.drawString(50, y, "ID | Название | Цена | Количество")
    y -= 20
    for flower in Flower.objects.all():
        p.drawString(50, y, f"{flower.id} | {flower.name} | {flower.price} | {flower.stock}")
        y -= 20

    # Закрываем PDF
    p.showPage()
    p.save()
    return response


def generate_docx(request):
    # Создаем документ
    doc = Document()
    doc.add_heading("Отчет по цветочному магазину", level=1)
    doc.add_paragraph(f"Дата создания: {datetime.now().date()}")
    doc.add_heading("Краткое описание сервиса", level=2)
    doc.add_paragraph("Сервис предоставляет REST API для работы с цветами, клиентами и заказами.")

    # Добавляем таблицу цветов
    doc.add_heading("Таблица цветов", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Название"
    hdr_cells[2].text = "Цена"
    hdr_cells[3].text = "Количество"

    for flower in Flower.objects.all():
        row_cells = table.add_row().cells
        row_cells[0].text = str(flower.id)
        row_cells[1].text = flower.name
        row_cells[2].text = str(flower.price)
        row_cells[3].text = str(flower.stock)

    # HTTP response для DOCX
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="flower_shop_report_{datetime.now().date()}.docx"'
    doc.save(response)
    return response